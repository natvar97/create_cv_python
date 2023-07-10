from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

document.add_picture('profile_picture.jpg',
                     width=Inches(2.0),
                     height=Inches(2.0))

# Add details about your self
name = input('What is your name? : ')
speak('Hello ' + name + ' How are you?')
speak('What is your mobile number?')
mobile = input('What is your mobile? : ')
email = input('What is your email address? : ')
paragraph_string = name + ' | ' + mobile + ' | ' + email
document.add_paragraph(paragraph_string)

# About me
document.add_heading('About me')
document.add_paragraph('Tell me about your self')

# Work experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter company name : ')
from_date = input('From date : ')
to_date = input('To date : ')

p.add_run(company + ' ').bold = True
p.add_run(', ' + from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ' : ')
p.add_run(experience_details)

# More experience details
while True:
    has_more_experience = input('Do you have more experiences? (y/n): ')
    if has_more_experience.lower() == 'y':
        p = document.add_paragraph()
        company = input('Enter company name : ')
        from_date = input('From date : ')
        to_date = input('To date : ')

        p.add_run(company + ' ').bold = True
        p.add_run(', ' + from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' : ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
skill = input('Enter your skill : ')
skills_p = document.add_paragraph(skill)
skills_p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you want to add your skills to your CV? (y/n): ')
    if has_more_skills.lower() == 'y':
        skill = input('Enter your skill : ')
        skills_p = document.add_paragraph(skill)
        skills_p.style = 'List Bullet'
    else:
        break

# Footer
sections = document.sections
if len(sections) > 0:
    section = sections[0]
    footer = section.footer
    paragraphs = footer.paragraphs
    if len(paragraphs) > 0:
        paragraph = paragraphs[0]
        paragraph.text = "CV generated using Indialone Artificial Intelligence code"

document.save("cv.docx")
